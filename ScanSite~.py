from bs4 import BeautifulSoup
import urllib.request
import winsound
import docx
import re
# import docx

# import time
from docx import Document
import datetime
from docx.shared import Inches,Pt,Mm,RGBColor
from docx.text.run import Font, Run
from docx.enum.text import WD_COLOR_INDEX,WD_ALIGN_PARAGRAPH
from docx.styles.style import _ParagraphStyle
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ROW_HEIGHT,WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell


# from docx.font.highlight_color import WD_COLOR_INDEX 
# import docx.shared.Length
def open_list(listOfSpels):
    f=open('listOfSpels.txt','r',encoding='utf-8')
    adding_flag=""
    list_of_keys=[]
    for i in range(1,21):
        list_of_keys.append(str(i))
    list_of_keys=list_of_keys+["к4","к6","к8","к10","к12","к20"]
    # print(list_of_keys)
    i=0
    for line in f:
        if line.strip()!='!!!':
            if line.strip()!="":
                if line.strip() in list_of_keys:
                    adding_flag=line.strip()+". "
                else:
                    listOfSpels[i].append(adding_flag+line.strip())
                    adding_flag=""
        else:
            listOfSpels.append([])
            i=i+1
    listOfSpels.remove([])
    f.close()
    return listOfSpels
def add_row(table0,t):
    if t%8==0:
        # print("add 5 rows")
        cur_row= table0.add_row()
        cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        cur_row.height=Inches(10/25.4*19.8/2)
        cur_row= table0.add_row()
        cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        cur_row.height=Inches(10/25.4*19.8/2)

        cur_row= table0.add_row()
        cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        cur_row.height=Inches(10/25.4*0.05)

        cur_row= table0.add_row()
        cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        cur_row.height=Inches(10/25.4*20/2)
        cur_row= table0.add_row()
        cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        cur_row.height=Inches(10/25.4*20/2)

        cur_row= table0.add_row()
        cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        cur_row.height=Inches(10/25.4*0.05)
def set_cell_border(cell: _Cell, **kwargs):
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV','left','right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
        
# httml='http://www.python.org/'

# Страница со списком заклинаний 
httml='https://dungeon.su/spells/'
with urllib.request.urlopen(httml) as f:
    listOfSitesOfSpells=[]
    soup = BeautifulSoup (f, 'html.parser')
    # Ищем все ссылки
    for link in soup.find_all('a'):
        # Ищем ссылки со '/spells/'
        if link.get('href').find('/spells/')>=0 and link.get('href').find('http')==-1 and len(link.get('href'))>len('/spells/'):
            # Ищем ссылки Spells в которых есть номер заклинания
            if link.get('href').find('1')!=-1 or link.get('href').find('2')!=-1 or link.get('href').find('3')!=-1 or link.get('href').find('4')!=-1 or link.get('href').find('5')!=-1 or link.get('href').find('6')!=-1 or link.get('href').find('7')!=-1 or link.get('href').find('8')!=-1 or link.get('href').find('9')!=-1:
                # print(link.get('href'))
                # Составляем список ссылок на заклинания
                listOfSitesOfSpells.append('https://dungeon.su'+link.get('href'))
listOfSpels=[[]]
kolOfSpels=len(listOfSitesOfSpells)
# kolOfSpels=10
i=0
# if 1:

# Открываем listOfSpels.txt в котором хранятся заклы, разделённые !!!
try:
    # Открываем файл, записываем в оперативку все заклы
    listOfSpels=open_list(listOfSpels)
    
    # print(listOfSpels)
    # print(len(listOfSpels))
# except Exception:
    # print(Exception)
except:
    print("No file")
    f=open('listOfSpels.txt','w',encoding='utf-8')
    f.close()

# Убираем пустые строки в конце
if len(listOfSpels)!=0: 
    while listOfSpels[len(listOfSpels)-1]=='':
            listOfSpels.remove[len(listOfSpels)-1]
# print(len(listOfSpels))
# print(listOfSpels)

# Если список заклинаний по длине не совпадает
if len(listOfSpels)!=kolOfSpels:
    print("Len not same")
    listOfSpels=[[]]
    for i in range(0,kolOfSpels):
    # if 1:
        # i=15
        # Добавляем заготовку для нового заклинания
        listOfSpels.append([])
        httml=listOfSitesOfSpells[i]
        with urllib.request.urlopen(httml) as f:
            # print(f.read().decode('utf-8'))
            html_doc= f
            soup = BeautifulSoup (html_doc, 'html.parser')
            if 0:
                print(soup.get_text())
            else:
                pos=soup.title.string.find('(')
                # print(pos)
                txt=soup.title.string[0:pos]
                while txt.rfind(' ')+1==len(txt):
                    pos=txt.rfind(' ')
                # print (str(pos)+"   "+str(len(txt)))
                    txt=soup.title.string[0:pos]
                title=txt
                
                txt=soup.get_text()
                # print(title+"!")
                pos0=txt.find(title)
                pos0=txt.find(title,pos0+1)
                pos0=txt.find(title,pos0+1)
                pos1=txt.find(']',pos0)+1
                spell_name=txt[pos0:pos1]
                # print("0!"+spell_name+"!")
                print(spell_name)
                listOfSpels[i].append(spell_name)

                pos0=pos1
                pos1=txt.find('Время накладыван',pos0)
                spell_lvl=txt[pos0:pos1]
                # print("1!"+spell_lvl+"!")
                listOfSpels[i].append(spell_lvl)

                pos0=pos1
                pos1=txt.find('Дистанция',pos0)
                spell_time=txt[pos0:pos1]
                # print("2!"+spell_time+"!")
                listOfSpels[i].append(spell_time)
                
                pos0=pos1
                pos1=txt.find('Компоненты',pos0)
                spell_dist=txt[pos0:pos1]
                # print("3!"+spell_dist+"!")
                listOfSpels[i].append(spell_dist)

                pos0=pos1
                pos1=txt.find('Длитель',pos0)
                spell_komponents=txt[pos0:pos1]
                # print("4!"+spell_komponents+"!")
                listOfSpels[i].append(spell_komponents)

                if txt.find('Классы',pos0)!=-1:
                    pos0=pos1
                    pos1=txt.find('Классы',pos0)
                    spell_dlitelnost=txt[pos0:pos1]
                    # print("5!"+spell_dlitelnost+"!")
                    listOfSpels[i].append(spell_dlitelnost)
                else:
                    listOfSpels[i].append("???")
                    # pos0=pos1
                    # pos1=txt.find('Архетипы',pos0)
                    # spell_dlitelnost=txt[pos0:pos1]
                    # # print("7!"+spell_arhetips+"!")
                    # listOfSpels[i].append(spell_dlitelnost)

                TakeArch=False
                postmp=pos0
                if txt.find('Архетипы',pos0)!=-1:
                    pos0=pos1
                    pos1=txt.find('Архетипы',pos0)
                    # print("!!!"+str(pos1)+"!!!")
                    spell_klasses=txt[pos0:pos1]
                    # print("6!"+spell_klasses+"!")
                    listOfSpels[i].append(spell_klasses)
                else:
                    pos0=pos1
                    pos1=txt.find('Источник',pos0)
                    spell_arhetips=txt[pos0:pos1]
                    # print("7!"+spell_arhetips+"!")
                    listOfSpels[i].append(spell_arhetips)
                    postmp=pos0+1
                    TakeArch=True

                if txt.find('Архети',postmp)!=-1 and not TakeArch:
                    pos0=txt.find('Архетипы',postmp)
                    pos1=txt.find('Источник',pos0)
                    spell_arhetips=txt[pos0:pos1]
                    # print("7!"+spell_arhetips+"!")
                    listOfSpels[i].append(spell_arhetips)
                else:
                    listOfSpels[i].append("???")
                    pass

                pos0=pos1
                # pos0=txt.find('»',pos0)+1
                # pos1=10000000
                # for cur_char in range(1040,1072):                    
                #     if txt.find(chr(cur_char),pos0)+1<pos1:
                #         pos1=txt.find(chr(cur_char),pos0)+1
                # spell_source=txt[pos0:pos1]
                # # print("8!"+spell_source+"!")
                # listOfSpels[i].append(spell_source)

                pos1=txt.find('»',pos0)+1
                if txt.find('»',pos1+1)+1-pos1<100 and txt.find('»',pos1+1)!=-1:
                    pos1=txt.find('»',pos1+1)+1
                spell_source=txt[pos0:pos1]
                # print("8!"+spell_source+"!")
                listOfSpels[i].append(spell_source)
                
                pos0=pos1
                pos1=txt.find('Поделиться',pos0)-1
                spell_text=txt[pos0:pos1]
                while spell_text.rfind(' ')+1==len(spell_text):
                        pos=spell_text.rfind(' ')
                        spell_text=spell_text[0:pos]
                # while spell_text.rfind('\n\n')!=-1:
                #     spell_text.replace('\n\n', '\n')
                spell_text=spell_text[0:len(spell_text)-3]
                # print("9!"+spell_text+"!")
                # spell_text.replace("\n\n","\n")
                # spell_text.replace("\r\r","\r")
                listOfSpels[i].append(spell_text)
                
                # print()
                # print(soup.title.name)
                # print(soup.prettify())
    try:
        while 1:
            listOfSpels.remove([])
    except:
        pass
    # print(listOfSpels)
    f = open('listOfSpels.txt','w',encoding='utf-8')
    for element in listOfSpels:
        for i in element:
            # print(str(i))
            f.write(str(i))
            f.write('\n')
        f.write('!!!\n')
    f.close()
    listOfSpels=[[]]
    listOfSpels=open_list(listOfSpels)
else:
    print("SAME!"+str(kolOfSpels))
    print()
winsound.Beep(500, 200)
# print(listOfSpels)
listOfClasses=['Бард','Варвар','Воин','Волшебник','Друид','Жрец','Изобретатель','Колдун','Монах','Паладин','Плут','Следопыт','Чародей','Все классы','Настройка']
fChoice=True
fFlagChoice=0
Main_Indent=3
while fChoice:
    choiceClass=15
    # while choiceClass>14:
    fChoice=not fChoice
    for i in range (0,len(listOfClasses)):
        print(str(i+1)+". "+listOfClasses[i])

    choiceClass=int(input("Выберите класс: "))
    
    if choiceClass==1: # 1. Бард

        pass
    elif choiceClass==2:# 2. Варвар
        print("У Варвара нет магии")
        pass
    elif choiceClass==3:# 3. Воин
        while fFlagChoice!=1 and fFlagChoice!=2:
            print("Вы мистический воин?\n1. Да\n2. Нет")
            fFlagChoice= int(input())
        pass
    elif choiceClass==4:# 4. Волшебник

        pass
    elif choiceClass==5:# 5. Друид

        pass
    elif choiceClass==6:# 6. Жрец

        pass
    elif choiceClass==7:# 7. Изобретатель

        pass
    elif choiceClass==8:# 8. Колдун

        pass
    elif choiceClass==9:# 9. Монах

        pass
    elif choiceClass==10:# 10. Паладин

        pass
    elif choiceClass==11:# 11. Плут

        pass
    elif choiceClass==12:# 12. Следопыт

        pass
    elif choiceClass==13:# 13. Чародей
        pass
    elif choiceClass==14:# 14. Все классы
        pass
    elif choiceClass==15:# 15. Настройки
        print("Настройка левой границы")
        Main_Indent=int(input("Размеры левого отступа(мм): "))
        # print(Main_Indent+'мм')
        fChoice=True
        pass        
    else:
        print("Нет такого класса")
        fChoice=True
# Main_Indent=13
# Main_Indent=3
for i in range (0,len(listOfClasses)):
            listOfClasses[i]=listOfClasses[i].lower()
listOfLvl=-1
while listOfLvl<0 or listOfLvl>10:
    listOfLvl=3
    listOfLvl=int(input("Уровень заклинания (0-9) или 10 для всех заклинаний: "))
if listOfLvl==10:
    Spell_lvl_s=0
    Spell_lvl_f=10
else:
    Spell_lvl_s=listOfLvl
    Spell_lvl_f=listOfLvl+1


# print(listOfClasses[int(choiceClass)-1])
kolClass=0
kolArhetip=0
# print("Spisok dlya zapisi")
while listOfSpels[len(listOfSpels)-1]=='':
        listOfSpels.remove('')
# print(listOfSpels)

doc = Document()
# doc.add_heading('Document Title', 0)
section = doc.sections[-1]
section.orientation = 1
# section.page_width = Inches(10/25.4*29.7)
section.page_width =Mm(297)
# section.page_height = Inches(10/25.4*21)
section.page_height = Mm(210) 
section.top_margin=Mm(5)
section.bottom_margin=Mm(5)
section.left_margin = Mm(7)
section.right_margin = Mm(7)

table0 = doc.add_table(rows=6, cols=4)
# modifyBorder(table0)
table0.autofit=False
hdr_cells = table0.rows[0].cells
table0.rows[0].height_rule = WD_ROW_HEIGHT.EXACTLY
table0.rows[0].height=Inches(10/25.4*19.8/2)
table0.rows[1].height_rule = WD_ROW_HEIGHT.EXACTLY
table0.rows[1].height=Inches(10/25.4*19.8/2)
table0.rows[2].height_rule = WD_ROW_HEIGHT.EXACTLY
table0.rows[2].height=Inches(10/25.4*0.05)
table0.rows[3].height_rule = WD_ROW_HEIGHT.EXACTLY
table0.rows[3].height=Inches(10/25.4*20/2)
table0.rows[4].height_rule = WD_ROW_HEIGHT.EXACTLY
table0.rows[4].height=Inches(10/25.4*20/2)
table0.rows[5].height_rule = WD_ROW_HEIGHT.EXACTLY
table0.rows[5].height=Inches(10/25.4*0.05)

hdr_cells[0].width=Inches(10/25.4*28.3/4)
hdr_cells[1].width=Inches(10/25.4*28.3/4)
hdr_cells[2].width=Inches(10/25.4*28.3/4)
hdr_cells[3].width=Inches(10/25.4*28.3/4)
t=0
row=0
column=0
back_name=listOfClasses[choiceClass-1]
print (back_name)

for listOfLvl in range (Spell_lvl_s,Spell_lvl_f):
    if listOfLvl==0:
        listOfLvl="Заговор"
    for i in listOfSpels:
        if i[1].find(str(listOfLvl))!=-1 or listOfLvl==10:
            # row=(t//4)
            # // - целочисленное деление; % - Остаток от деления
            # row=t//4+t//8*5
            
            
            row=(t//8*6)+t//4%2
            # print()
            column=t%4
            hdr_cells = table0.rows[row].cells
            cell = table0.cell(row, column)
            set_cell_border(
                cell,
                left={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                right={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                top={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                bottom={"sz": 1, "val": "single", "color": "#000000", "space": "0"}

            )
            
            if i[6].find(listOfClasses[3])!=-1 and choiceClass==3 and fFlagChoice==1 :#Мистический рыцарь and (i[1].find('воплощен')!=-1 or i[1].find('огражден')!=-1)
                choiceClass=4
            if i[6].find(listOfClasses[int(choiceClass)-1])!=-1 or i[7].find(listOfClasses[int(choiceClass)-1])!=-1 or choiceClass==14:
                font_size0=8
                font_size1=8
                chars_in_line=39
                lenn=0
                cur_String=0
                cur_String_len=0
                for j in range(7,len(i)):
                    lenn=lenn+len(i[j].replace(" ",""))
                    pass
                print(i[0]+" "+str(lenn))
                big_text=False
                # if lenn>2000:
                #     font_size0=6
                #     font_size1=6
                #     big_text=True
                # elif lenn>1800:
                #     big_text=True
                # elif lenn>1460:
                #     font_size0=5
                #     font_size1=5
                # elif lenn>1264:
                #     font_size0=5
                
                #     font_size1=6
                # # elif lenn>935:
                # #     font_size0=6
                # #     font_size1=6
                # elif lenn>834:
                #     font_size0=6
                #     font_size1=6
                # elif lenn>583:
                #     font_size0=7
                #     font_size1=7
                # elif lenn>536:
                #     font_size0=7
                #     font_size1=8
            
                cell.paragraphs[0].paragraph_format.space_after=Mm(0)
                cell.paragraphs[0].paragraph_format.line_spacing=1
                cell.paragraphs[0].paragraph_format.left_indent=Mm(Main_Indent)
                cell.paragraphs[0].paragraph_format.space_before=Mm(1)
                # cell.text='!!!'
                # print(cell.paragraphs)
                cur_run=cell.paragraphs[-1].add_run(i[0])
                cur_run.font.name='TeXGyreCursor'
                cur_run.bold=True
                cur_run.font.size=Pt(8)
                cur_run.font.color.rgb=RGBColor(255,255,255)
                cur_run.font.highlight_color = WD_COLOR_INDEX.BLACK
                # cell.add_run('&&&')

                
                # styles = doc.styles
                # style=styles.add_style('Main1',WD_STYLE_TYPE.PARAGRAPH)
                
                # cur_run.bold = True
                # paragraph.add_run('bold').bold = True
                # par_style=_ParagraphStyle
                correct=0
                start_txt=9
                # if i[7]=="???":
                #     start_txt=8
                # print(i[2])
                correcting=True
                while correcting:
                    for cur_par in range(1,len(i)):
                        if i[cur_par]!="???":
                            cur_String+=1
                            cur_String_len=0
                            splitted=re.split(' ', i[cur_par])
                            for xi in range(len(splitted)):
                                    splitted[xi]=splitted[xi]+" "
                                    # cur_run=paragraph.add_run(splitted[xi])
                                    if cur_String_len+len(splitted[xi])>chars_in_line:
                                        cur_String_len=len(splitted[xi])
                                        cur_String+=1
                                        # print(splitted[xi])
                                        
                                    else:
                                        cur_String_len+=len(splitted[xi])
                    if cur_String >65-5*font_size0 and font_size0!=5:
                        font_size0-=1
                        font_size1-=1
                        if font_size0==7:
                            chars_in_line=43
                        elif font_size0==6:
                            chars_in_line=50
                        elif font_size0==5:
                            chars_in_line=61
                        print(cur_String)
                        cur_String=0
                        cur_String_len=0

                    else:
                        correcting=False
                        print(cur_String)

                paragraph = cell.add_paragraph('')
                cell.paragraphs[1].paragraph_format.space_after=Mm(0)
                cell.paragraphs[1].paragraph_format.line_spacing=1
                cell.paragraphs[1].paragraph_format.left_indent=Mm(Main_Indent)
                cur_run=paragraph.add_run(i[1])
                cur_run.font.name='TeXGyreCursor'
                cur_run.font.size=Pt(font_size0)
                cur_run.italic=True
                # Пишем на карточку описание 
                for cur_par in range(2,start_txt):
                    if i[cur_par]!="???":
                            # splitted=re.split(' ', i[cur_par])
                            paragraph = cell.add_paragraph('')
                            cell.paragraphs[cur_par-correct].paragraph_format.space_after=Mm(0)
                            cell.paragraphs[cur_par-correct].paragraph_format.line_spacing=1
                            cell.paragraphs[cur_par-correct].paragraph_format.left_indent=Mm(Main_Indent)
                            cur_run_dvoetoch=i[cur_par].find(":")
                            
                            # if cur_run_dvoetoch!=-1:
                            


                            cur_run=paragraph.add_run(i[cur_par][0:cur_run_dvoetoch+1])
                            cur_run.font.name='TeXGyreCursor'
                            cur_run.bold=True
                            cur_run.font.size=Pt(font_size0)
                            cur_run=paragraph.add_run(i[cur_par][cur_run_dvoetoch+1:len(i[cur_par])])
                            cur_run.font.name='TeXGyreCursor'
                            cur_run.font.size=Pt(font_size0)

                    else:
                        # print(cur_par)
                        correct=1
                
                # Пишем на карточку текст описания заклинания
                if not big_text:
                    for cur_par in range(start_txt,len(i)):
                        # splitted=i[cur_par].split()
                        
                        paragraph = cell.add_paragraph('')
                        cell.paragraphs[cur_par-correct].paragraph_format.space_after=Mm(0)
                        cell.paragraphs[cur_par-correct].paragraph_format.line_spacing=1
                        cell.paragraphs[cur_par-correct].paragraph_format.left_indent=Mm(Main_Indent)
                        cell.paragraphs[cur_par-correct].paragraph_format.first_line_indent=Mm(1)
                        # print(splitted)
                        # print()
                        
                            
                            
                            
                        if i[cur_par].find("На больших уровнях.")!=-1:
                            # print("На больших уровнях.")
                            cur_run=paragraph.add_run("На больших уровнях.")
                            cur_run.font.name='TeXGyreCursor'
                            cur_run.font.size=Pt(font_size1)
                            cur_run.bold=True
                            i[cur_par]=i[cur_par][i[cur_par].find("На больших уровнях.")+len("На больших уровнях."):len(i[cur_par])]
                        if i[cur_par].find("Компонент авторских отчислений (А)")!=-1:
                            # print("Авторские отчисления")
                            cur_run=paragraph.add_run("Компонент авторских отчислений (А) ")
                            cur_run.font.name='TeXGyreCursor'
                            cur_run.font.size=Pt(font_size1)
                            cur_run.italic=True
                            cur_run.bold=True
                            i[cur_par]=i[cur_par][i[cur_par].find("Компонент авторских отчислений (А)")+len("Компонент авторских отчислений (А)"):len(i[cur_par])]
                        
                        cur_run=paragraph.add_run(i[cur_par])
                        cur_run.font.name='TeXGyreCursor'
                        cur_run.font.size=Pt(font_size1)
                else:
                    
                    lenn=0
                    for cur_par in range(start_txt,len(i)):
                        # if i[cur_par]!="???":

                        if i[cur_par].find("На больших уровнях.")!=-1:
                                # print("На больших уровнях.")
                                cur_run=paragraph.add_run("На больших уровнях.")
                                cur_run.font.name='TeXGyreCursor'
                                cur_run.font.size=Pt(font_size1)
                                cur_run.bold=True
                                i[cur_par]=i[cur_par][i[cur_par].find("На больших уровнях.")+len("На больших уровнях."):len(i[cur_par])]
                        if i[cur_par].find("Компонент авторских отчислений (А)")!=-1:
                                # print("Авторские отчисления")
                                cur_run=paragraph.add_run("Компонент авторских отчислений (А) ")
                                cur_run.font.name='TeXGyreCursor'
                                cur_run.font.size=Pt(font_size1)
                                cur_run.italic=True
                                cur_run.bold=True
                                i[cur_par]=i[cur_par][i[cur_par].find("Компонент авторских отчислений (А)")+len("Компонент авторских отчислений (А)"):len(i[cur_par])]
                        lenn=lenn+len(i[cur_par])
                        if lenn<834:
                            paragraph = cell.add_paragraph('')
                            cell.paragraphs[cur_par-correct].paragraph_format.space_after=Mm(0)
                            cell.paragraphs[cur_par-correct].paragraph_format.line_spacing=1
                            cell.paragraphs[cur_par-correct].paragraph_format.left_indent=Mm(Main_Indent)
                            cur_run=paragraph.add_run(i[cur_par])
                            cur_run.font.name='TeXGyreCursor'
                            cur_run.font.size=Pt(font_size1)
                        else:
                            print(cur_par)
                            start_txt=cur_par
                            break

                cell= table0.cell(row+3, 3-column)
                # cell.paragraphs[0].paragraph_format.right_indent=Mm(Main_Indent)
                # cell.paragraphs[0].paragraph_format.left_indent=Mm(8)
# ------------------------------------------------------------------------------------------------------------------------------------------------------------------------
                # Если есть большой текст, то вторая сторона для продолжения текста
                if big_text:
                    big_text=False
                    cur_par_ind=0
                    
                    # paragraph = cell.add_paragraph('')
                    # space_before
                    # cur_run=paragraph.add_run('')
                    # cur_run.font.name='TeXGyreCursor'
                    # cur_run.font.size=Pt(3)
                    for cur_par in range(start_txt,len(i)):
                        cell.paragraphs[cur_par_ind].paragraph_format.space_after=Mm(0)
                        cell.paragraphs[cur_par_ind].paragraph_format.space_before=Mm(0)
                        cell.paragraphs[cur_par_ind].paragraph_format.line_spacing=1
                        cell.paragraphs[cur_par_ind].paragraph_format.right_indent=Mm(Main_Indent)
                        cell.paragraphs[cur_par_ind].paragraph_format.left_indent=Mm(5)
                        paragraph=cell.paragraphs[cur_par_ind]
                        cur_par_ind=cur_par_ind+1
                        # На больших уровнях. 
                        if i[cur_par].find("На больших уровнях.")!=-1:
                                # print("На больших уровнях.")
                                cur_run=paragraph.add_run("На больших уровнях.")
                                cur_run.font.name='TeXGyreCursor'
                                cur_run.font.size=Pt(font_size1)
                                cur_run.bold=True
                                i[cur_par]=i[cur_par][i[cur_par].find("На больших уровнях.")+len("На больших уровнях."):len(i[cur_par])]
                        if i[cur_par].find("Компонент авторских отчислений (А)")!=-1:
                                # print("Авторские отчисления")
                                cur_run=paragraph.add_run("Компонент авторских отчислений (А) ")
                                cur_run.font.name='TeXGyreCursor'
                                cur_run.font.size=Pt(font_size1)
                                cur_run.italic=True
                                cur_run.bold=True
                                i[cur_par]=i[cur_par][i[cur_par].find("Компонент авторских отчислений (А)")+len("Компонент авторских отчислений (А)"):len(i[cur_par])]
                        cur_run=paragraph.add_run(i[cur_par])
                        cur_run.font.name='TeXGyreCursor'
                        cur_run.font.size=Pt(font_size1)
                        paragraph = cell.add_paragraph('')
                    cell.paragraphs[0].paragraph_format.space_before=Mm(1)
                    pass

                # Иначе большими буквами название заклинания, круг и классы
                else:        
                    
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.paragraphs[0].paragraph_format.right_indent=Mm(Main_Indent)
                    cell.paragraphs[0].paragraph_format.left_indent=Mm(5)
                    txt=i[0]
                    txt=txt[0:txt.find("[")-1]
                    cur_run=cell.paragraphs[0].add_run(txt)
                    # cur_run.font.name='TeXGyreCursor'
                    cur_run.font.size=Pt(16)
                    cur_run.bold=True
                    paragraph = cell.add_paragraph('')
                    paragraph.paragraph_format.right_indent=Mm(Main_Indent)
                    paragraph.paragraph_format.left_indent=Mm(5)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if listOfLvl==0 or listOfLvl=="Заговор":
                        cur_run=paragraph.add_run("Заговор")
                        # cur_run.font.name='TeXGyreCursor'
                    else:

                        cur_run=paragraph.add_run(str(listOfLvl) + " круг")
                        # cur_run.font.name='TeXGyreCursor'
                    cur_run.font.size=Pt(13)
                    txt=i[6]
                    txt=txt[txt.find(":")+2:len(txt)]
                    txt=txt.title()
                    paragraph = cell.add_paragraph('')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.right_indent=Mm(Main_Indent)
                    paragraph.paragraph_format.left_indent=Mm(5)
                    cur_run=paragraph.add_run(txt)
                    cur_run.font.size=Pt(13)
                    if i[7]!="???":
                        txt=i[7]
                        txt=txt[txt.find(":")+2:len(txt)]
                        # paragraph = cell.add_paragraph('')
                        # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cur_run=paragraph.add_run(", "+txt)
                        # cur_run.font.name='TeXGyreCursor'
                        cur_run.font.size=Pt(13)
                    # cur_run.font.all_caps=True

                
                # if len(i)==100:
                #     # print(i[9])
                #     paragraph = cell.add_paragraph('')
                #     cell.paragraphs[9].paragraph_format.space_after=Mm(0)
                #     cell.paragraphs[9].paragraph_format.line_spacing=1
                #     cur_run=paragraph.add_run(i[9])
                #     cur_run.font.size=Pt(8)
                # except:
                #     pass
                # paragraph.add_run(i[2]).italic=True            
                # par_style = styles.add_style("Par_Name", WD_STYLE_TYPE.PARAGRAPH)
                # par_font=Font
                # par_font.color=RGBColor(0xB5,0x0C,0xB2)
                # par_font.size=Pt(8)
                # print(par_font.size)
                # styles = doc.styles
                # par_style.font=par_font
                # par_style.name="Main1"
                # 
                # style_font=style.font
                # style_font.size=Pt(8)
                # style_font.color=RGBColor(0xB5,0x0C,0xB2)
                # style.font=style_font
                # style_font.color=
                # print(style.font.size)
                # k=0
                # for stt in styles:
                #     if stt.name=="Main1":
                #         print(stt.font.size)
                #         # stt.font=par_font
                #         k=k+1
                # print("k="+str(k))
                # par_style.name="Par_Name"
                # cell.paragraphs[0].style='Main1'
                # print(par_font)
                # RGBColor(0xB5,0x0C,0xB2)
                # cell.paragraphs[0].style.font.color=RGBColor(0xB5,0x0C,0xB2)
                # print(cell.paragraphs[0].text)

                # print(cell.paragraphs[0])
                # cell.add_table(2,2)
                # cell.run.bold=True
                # par_for=doc.styles['Normal'].paragraph_format
                # 
                # paragraph.space_before=Pt(1)
                # paragraph.space_after=Pt(1)
                # all_paras=table0.paragraphs
                # par_for=paragraph.paragraph_format.space_after=Pt(0)
                # prior_paragraph = paragraph.insert_paragraph_before(i[0])
                # cell.font.size=Pt(8)
                # cell.font.highlight_color=(0,0,0)
                # hdr_cells[column].add_paragraph(i[0], style='List Number')
                # print(i[1])
                # print(i[2])
                # print(i[3])
                # print(i[4])
                # print(i[5])
                # print(i[6])
                if i[6].find(listOfClasses[int(choiceClass)-1])!=-1:
                    kolClass+=1
                if i[7].find(listOfClasses[int(choiceClass)-1])!=-1:
                    # print(i[7])
                    kolArhetip+=1
                # print(i[8])
                
                # print()
                # toDoc(i,doc,t)
                # print ("t="+str(t)+"  row="+str(row))
                t=t+1
                add_row(table0,t)
            # if t>=16:
            #     break
            

                
                
            
        
    
print("Klass: "+str(kolClass)+"  Arhetips: "+str(kolArhetip))
name =listOfClasses[int(choiceClass)-1]+" "+str(listOfLvl)+"_lvl "+ str(datetime.datetime.now().time())
name=name.replace(':','_',2)
doc.save('Docs/'+name+'.docx')