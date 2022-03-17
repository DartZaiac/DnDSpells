# Необходимые библиотеки
# pip install beautifulsoup4
# pip install docx2pdf
# pip install python-docx

from bs4 import BeautifulSoup
import urllib.request
import docx
import re

from docx import Document
import datetime
from docx.shared import Inches,Pt,Mm,RGBColor
from docx.text.run import Font, Run
from docx.enum.text import WD_COLOR_INDEX,WD_ALIGN_PARAGRAPH
from docx.styles.style import _ParagraphStyle
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ROW_HEIGHT,WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

from docx2pdf import convert

import os, ssl
if (not os.environ.get('PYTHONHTTPSVERIFY', '') and getattr(ssl, '_create_unverified_context', None)):
    ssl._create_default_https_context = ssl._create_unverified_context


# from docx.font.highlight_color import WD_COLOR_INDEX 
# import docx.shared.Length
def open_list(listOfSpels):
    f=open('listOfSpells.txt','r',encoding='utf-8')
    adding_flag=""
    list_of_keys=[]
    for i in range(1,21):
        list_of_keys.append(str(i))
    list_of_keys=list_of_keys+["к4","к6","к8","к10","к12","к20"]
    # print(list_of_keys)
    i=0
    for line in f:
        if line.strip()!='!!!':
            if line.strip()!="":
                if line.strip() in list_of_keys:
                    adding_flag=line.strip()+". "
                else:
                    listOfSpels[i].append(adding_flag+line.strip())
                    adding_flag=""
        else:
            listOfSpels.append([])
            i=i+1
    listOfSpels.remove([])
    f.close()
    return listOfSpels
def add_row(table0,t):
    if (t)%(columns_in_table*rows_in_table)==0:
        # print("add 5 rows")
        for adding_rows in range(rows_in_table):
            cur_row= table0.add_row()
            cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
            cur_row.height=Inches(10/25.4*intTableHeight/rows_in_table)
        # cur_row= table0.add_row()
        # cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        # cur_row.height=Inches(10/25.4*intTableHeight/rows_in_table)

        cur_row= table0.add_row()
        cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        cur_row.height=Inches(10/25.4*0.05)

        for adding_rows in range(rows_in_table):
            cur_row= table0.add_row()
            cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
            cur_row.height=Inches(10/25.4*(intTableHeight+0.2)/rows_in_table)
        # cur_row= table0.add_row()
        # cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        # cur_row.height=Inches(10/25.4*(intTableHeight+0.2)/rows_in_table)

        cur_row= table0.add_row()
        cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        cur_row.height=Inches(10/25.4*0.05)
def set_cell_border(cell: _Cell, **kwargs):
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV','left','right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
def jump_to_cell(cell):
    # print(cur_String)
    # print("Jump to another cell")
    cell.paragraphs[0].paragraph_format.right_indent=Mm(0)
    cell.paragraphs[0].paragraph_format.left_indent=Mm(Main_Indent)
    for cur_par in range(1,len(cell.paragraphs)):
        # print(cell.paragraphs[cur_par].text)
        cell.paragraphs[cur_par].paragraph_format.space_after=Mm(0)
        cell.paragraphs[cur_par].paragraph_format.line_spacing=1
        cell.paragraphs[cur_par].paragraph_format.left_indent=Mm(Main_Indent)
        
        if cell.paragraphs[cur_par].text.find(":")<21 and cell.paragraphs[cur_par].text.find(":")>-1 or cur_par<4:
            cell.paragraphs[cur_par].paragraph_format.first_line_indent=Mm(0)
            a=cell.paragraphs[cur_par].text.find(":")
            if a>-1:
                for ab in range(0,len(cell.paragraphs[cur_par].runs)):
                    cell.paragraphs[cur_par].runs[ab].bold=True
                    if cell.paragraphs[cur_par].runs[ab].text.find(":")!=-1:
                        break
                pass
        else:
            cell.paragraphs[cur_par].paragraph_format.first_line_indent=Mm(1)
       
        if cell.paragraphs[cur_par].text.find("На больших уровнях")!=-1:
            
                cell.paragraphs[cur_par].runs[0].bold=True 
                cell.paragraphs[cur_par].runs[1].bold=True 
                cell.paragraphs[cur_par].runs[2].bold=True 
            # cell.paragraphs[cur_par].add_run('')
            # for ab in range (len(cell.paragraphs[cur_par].runs)-1,-2,-1):
            #     cell.paragraphs[cur_par].runs[ab]=cell.paragraphs[cur_par].runs[ab-1]
            # cell.paragraphs[cur_par].runs[0].text="На больших уровнях."
            # cell.paragraphs[cur_par].runs[0].font.size=Pt(font_size0)
            # cell.paragraphs[cur_par].runs[0].bold=True
        #     pass

        if cell.paragraphs[cur_par].text.find("На высоких уровнях")!=-1:
            cell.paragraphs[cur_par].add_run('')
            for ab in range (len(cell.paragraphs[cur_par].runs)-1,0,-1):
                cell.paragraphs[cur_par].runs[ab]=cell.paragraphs[cur_par].runs[ab-1]
            cell.paragraphs[cur_par].runs[0].text="На высоких уровнях."
            cell.paragraphs[cur_par].runs[0].font.size=Pt(font_size0)
            cell.paragraphs[cur_par].runs[0].bold=True
            pass
            # for ab in range(0,len(cell.paragraphs[cur_par].runs)):
            #     cell.paragraphs[cur_par].runs[ab].bold=True
            #     if cell.paragraphs[cur_par].runs[ab].text.find("уровнях")!=-1:
            #         break
        if cell.paragraphs[cur_par].text.find("Компонент авторских отчислений ")!=-1:
            for ab in range(0,len(cell.paragraphs[cur_par].runs)):
                cell.paragraphs[cur_par].runs[ab].bold=True
                cell.paragraphs[cur_par].runs[ab].italic=True
                if cell.paragraphs[cur_par].runs[ab].text.find("(А)")!=-1:
                    ba=cell.paragraphs[cur_par].runs[ab].text.find("(А)")
                    cell.paragraphs[cur_par].runs[ab].text=cell.paragraphs[cur_par].runs[ab].text[3:]
                    cell.paragraphs[cur_par].runs[ab].bold=False
                    cell.paragraphs[cur_par].runs[ab].italic=False
                    # print(cell.paragraphs[cur_par].runs[ab].text)
                    cell.paragraphs[cur_par].runs[ab-1].text=cell.paragraphs[cur_par].runs[ab-1].text+"(А) "
                    break
    cell= table0.cell(row+3, 3-column)
    # cell.paragraphs[0].paragraph_format.right_indent=Mm(6)
    cell.paragraphs[0].paragraph_format.right_indent=Mm(Main_Indent)
    cell.paragraphs[0].paragraph_format.left_indent=Mm(6)
    cell.paragraphs[0].paragraph_format.line_spacing=1
    cell.paragraphs[0].paragraph_format.space_after=Mm(0)
    # cell.paragraphs[0].add_run=""
    cur_run=cell.paragraphs[-1].add_run(" ")
    cur_run.font.size=Pt(1)

    return cell
# httml='http://www.python.org/'

# Ширина на 1 символ
# Ширина 8   67.3/42= 1.6 mm = 0.16 cm  высота 3.3 mm = 0.33 cm = (8*5-7)/100
# Ширина 7   58.8/42= 1.4 mm = 0.14 cm  высота 3.3 mm = 0.28 cm = (7*5-7)/100
# Ширина 6   50.4/42= 1.2 mm = 0.12 cm  высота 3.3 mm = 0.23 cm = (6*5-7)/100
# Ширина 5   42.0/42= 1   mm = 0.10 cm  высота 3.3 mm = 0.18 cm = (5*5-7)/100

columns_in_table = 2
rows_in_table = 1
int_Otstup=0.65 # размер полей
intTableHeight=19.6  # cm высота места для поля по вертикали
intTableWidth=28.3

# Страница со списком заклинаний 
httml='https://dungeon.su/spells/'
with urllib.request.urlopen(httml) as f:
    listOfSitesOfSpells=[]
    soup = BeautifulSoup (f, 'html.parser')
    # Ищем все ссылки
    for link in soup.find_all('a'):
        # Ищем ссылки со '/spells/'
        if link.get('href').find('/spells/')>=0 and link.get('href').find('http')==-1 and len(link.get('href'))>len('/spells/'):
            # Ищем ссылки Spells в которых есть номер заклинания
            if link.get('href').find('1')!=-1 or link.get('href').find('2')!=-1 or link.get('href').find('3')!=-1 or link.get('href').find('4')!=-1 or link.get('href').find('5')!=-1 or link.get('href').find('6')!=-1 or link.get('href').find('7')!=-1 or link.get('href').find('8')!=-1 or link.get('href').find('9')!=-1:
                # print(link.get('href'))
                # Составляем список ссылок на заклинания
                listOfSitesOfSpells.append('https://dungeon.su'+link.get('href'))
listOfSpels=[[]]
kolOfSpels=len(listOfSitesOfSpells)
# kolOfSpels=10

options_standart=['3','4','2']
options=[]
i=0
# Открываем options.txt в котором хранятся настройки
try:
    # Открываем файл, записываем в оперативку все заклы
    f=open('options.txt','r',encoding='utf-8')
    for line in f:
        options.append(line)
        i+=1
    f.close()
except:
    print("No option file")
    f=open('options.txt','w',encoding='utf-8')
    f.close()
if len(options_standart)!=len(options):
    options=options_standart.copy()
    print("stnd="+str(len(options_standart))+"   opt="+str(len(options)))

Main_Indent=int(options[0])
print("Левый отступ равен "+str(Main_Indent))
columns_in_table=int(options[1])
print("Количество столбиков "+str(columns_in_table))
rows_in_table=int(options[2])
print("Количество строк "+str(rows_in_table))



f=open('options.txt','w',encoding='utf-8')
for element in options:
    f.write(str(element)+'\n')
        
f.close()


# Открываем listOfSpells.txt в котором хранятся заклы, разделённые !!!
try:
    # Открываем файл, записываем в оперативку все заклы
    listOfSpels=open_list(listOfSpels)
except:
    print("No file")
    print("База данных с заклинаниями отсутствует. Качаем заклинания с Dungeon.su и формируем БД. Это может занять несколько минут")
    f=open('listOfSpells.txt','w',encoding='utf-8')
    f.close()



# Убираем пустые строки в конце
if len(listOfSpels)!=0: 
    while listOfSpels[len(listOfSpels)-1]=='':
            listOfSpels.remove[len(listOfSpels)-1]
# print(len(listOfSpels))
# print(listOfSpels)

# Если список заклинаний по длине не совпадает
if len(listOfSpels)!=kolOfSpels:
    print("Len not same")
    listOfSpels=[[]]
    for i in range(0,kolOfSpels):
    # if 1:
        # i=15
        # Добавляем заготовку для нового заклинания
        listOfSpels.append([])
        httml=listOfSitesOfSpells[i]
        with urllib.request.urlopen(httml) as f:
            # print(f.read().decode('utf-8'))
            html_doc= f
            soup = BeautifulSoup (html_doc, 'html.parser')
            if 0:
                # print(soup.get_text()),
                pass
            else:
                pos=soup.title.string.find('(')
                # print(pos)
                txt=soup.title.string[0:pos]
                while txt.rfind(' ')+1==len(txt):
                    pos=txt.rfind(' ')
                # print (str(pos)+"   "+str(len(txt)))
                    txt=soup.title.string[0:pos]
                title=txt
                
                txt=soup.get_text()
                # print(title+"!")
                pos0=txt.find(title)
                pos0=txt.find(title,pos0+1)
                pos0=txt.find(title,pos0+1)
                # pos1=txt.find(']',pos0)+1
                pos1=txt.find(' [',pos0)+1
                spell_name=txt[pos0:pos1]
                # print("0!"+spell_name+"!")
                # print(spell_name)
                listOfSpels[i].append(spell_name)

                pos0= txt.find(']',pos1)+1
                pos1=txt.find('Время накладыван',pos0)
                spell_lvl=txt[pos0:pos1]
                # print("1!"+spell_lvl+"!")
                listOfSpels[i].append(spell_lvl)

                pos0=pos1
                pos1=txt.find('Дистанция',pos0)
                spell_time=txt[pos0:pos1]
                # print("2!"+spell_time+"!")
                listOfSpels[i].append(spell_time)
                
                pos0=pos1
                pos1=txt.find('Компоненты',pos0)
                spell_dist=txt[pos0:pos1]
                # print("3!"+spell_dist+"!")
                listOfSpels[i].append(spell_dist)

                pos0=pos1
                pos1=txt.find('Длитель',pos0)
                spell_komponents=txt[pos0:pos1]
                # print("4!"+spell_komponents+"!")
                listOfSpels[i].append(spell_komponents)

                if txt.find('Классы',pos0)!=-1:
                    pos0=pos1
                    pos1=txt.find('Классы',pos0)
                    spell_dlitelnost=txt[pos0:pos1]
                    # print("5!"+spell_dlitelnost+"!")
                    listOfSpels[i].append(spell_dlitelnost)
                else:
                    listOfSpels[i].append("???")
                    # pos0=pos1
                    # pos1=txt.find('Архетипы',pos0)
                    # spell_dlitelnost=txt[pos0:pos1]
                    # # print("7!"+spell_arhetips+"!")
                    # listOfSpels[i].append(spell_dlitelnost)

                TakeArch=False
                postmp=pos0
                if txt.find('Архетипы',pos0)!=-1:
                    pos0=pos1
                    pos1=txt.find('Архетипы',pos0)
                    # print("!!!"+str(pos1)+"!!!")
                    spell_klasses=txt[pos0:pos1]
                    # print("6!"+spell_klasses+"!")
                    listOfSpels[i].append(spell_klasses)
                else:
                    pos0=pos1
                    pos1=txt.find('Источник',pos0)
                    spell_arhetips=txt[pos0:pos1]
                    # print("7!"+spell_arhetips+"!")
                    listOfSpels[i].append(spell_arhetips)
                    postmp=pos0+1
                    TakeArch=True

                if txt.find('Архети',postmp)!=-1 and not TakeArch:
                    pos0=txt.find('Архетипы',postmp)
                    pos1=txt.find('Источник',pos0)
                    spell_arhetips=txt[pos0:pos1]
                    # print("7!"+spell_arhetips+"!")
                    listOfSpels[i].append(spell_arhetips)
                else:
                    listOfSpels[i].append("???")
                    pass

                pos0=pos1
                # pos0=txt.find('»',pos0)+1
                # pos1=10000000
                # for cur_char in range(1040,1072):                    
                #     if txt.find(chr(cur_char),pos0)+1<pos1:
                #         pos1=txt.find(chr(cur_char),pos0)+1
                # spell_source=txt[pos0:pos1]
                # # print("8!"+spell_source+"!")
                # listOfSpels[i].append(spell_source)

                pos1=txt.find('»',pos0)+1
                if txt.find('»',pos1+1)+1-pos1<100 and txt.find('»',pos1+1)!=-1:
                    pos1=txt.find('»',pos1+1)+1
                spell_source=txt[pos0:pos1]
                # print("8!"+spell_source+"!")
                listOfSpels[i].append(spell_source)
                
                pos0=pos1
                pos1=txt.find('Поделиться',pos0)-1
                spell_text=txt[pos0:pos1]
                while spell_text.rfind(' ')+1==len(spell_text):
                        pos=spell_text.rfind(' ')
                        spell_text=spell_text[0:pos]
                # while spell_text.rfind('\n\n')!=-1:
                #     spell_text.replace('\n\n', '\n')
                spell_text=spell_text[0:len(spell_text)-3]
                # print("9!"+spell_text+"!")
                # spell_text.replace("\n\n","\n")
                # spell_text.replace("\r\r","\r")
                listOfSpels[i].append(spell_text)
                
                # print()
                # print(soup.title.name)
                # print(soup.prettify())
    try:
        while 1:
            listOfSpels.remove([])
    except:
        pass
    # print(listOfSpels)
    f = open('listOfSpells.txt','w',encoding='utf-8')
    for element in listOfSpels:
        for i in element:
            # print(str(i))
            f.write(str(i))
            f.write('\n')
        f.write('!!!\n')
    f.close()
    listOfSpels=[[]]
    listOfSpels=open_list(listOfSpels)
else:
    print("SAME!"+str(kolOfSpels))
    print()
# print(listOfSpels)
listOfClasses=['Бард','Варвар','Воин','Волшебник','Друид','Жрец','Изобретатель','Колдун','Монах','Паладин','Плут','Следопыт','Чародей','Все классы','Настройка']
fChoice=True
fFlagChoice=0
# Main_Indent=3
TitleCard=[['','',''],['','',''],['','',''],['','',''],['','',''],['','',''],['','',''],['','',''],['','',''],['','','']]
TitleFont='Cambria (Body)'
while fChoice:
    try:
        choiceClass=15
        # while choiceClass>14:
        fChoice=not fChoice
        for i in range (0,len(listOfClasses)):
            print(str(i+1)+". "+listOfClasses[i])

        choiceClass=int(input("Выберите класс: "))
        
        if choiceClass==1: # 1. Бард

            pass
        elif choiceClass==2:# 2. Варвар
            print("У Варвара нет магии")
            pass
        elif choiceClass==3:# 3. Воин
            while fFlagChoice!=1 and fFlagChoice!=2:
                print("Вы мистический воин?\n1. Да\n2. Нет")
                fFlagChoice= int(input())
            pass
        elif choiceClass==4:# 4. Волшебник
            TitleFont = 'Fortuna Gothic FlorishC'
            TitleCard= [['Волшебник','Заговоры','Молния спрыгнула с руки волшебника и ничего не ожидающий воин в латных доспехах, который подбежал к нему получил сильнейший удар током.'],
                        ['Волшебник', '1 круг','Ты бронирован от пяток до макушки? А уклонишся ли ты от магических стрел?'],
                        ['Волшебник','2 круг','Воин подбежал к волшебнику и со всех сил ударил по нему. Но промахнулся. А маг, помахав ручкой расстворился и появился на краю скалы.'],
                        ['Волшебник','3 круг','Огненный шар решает любые проблемы.'],
                        ['Волшебник', '4 круг','- Кто сказал что я не кузнец?\n- Ты сам.\n- Я оговорился. Держи свой новый доспех!'],	
                        ['Волшебник', '5 круг','- А ты точно не будешь наши мысли читать?\n- Нет это другого рода связь.\n- Ладно, соединяй.\n- Всем привет. Спарки может говорить! А вы знали как весело гоняться за своим хвостом?' ],
                        ['Волшебник','6 круг','-Как мы переберёмся на другую сторону? Там же 300 футов минимум! Ни один живой мост не достанет!\n- Магические врата ещё никто не отменял. Так что я, наверное, их построю.'],
                        ['Волшебник','7 круг','Ты предлагешь МНЕ отдыхать в этом убогом домишке? Ну нет, я не для того учил магические искусства 20 лет, чтобы отказать себе отдохнуть в приличном, хоть и магическом особняке.'],
                        ['Волшебник','8 круг','- Привет! Куда тебя нелёгкая занесла?\n- Сколько лет, сколько зим! Я тебя даже незаметил.\n- А я могу из любой точки плана подключиться.\n- Жуть, ты как это сделал? Я сейчас в Невервинтере!\n- Я тоже! Пошли в кабак!'],
                        ['Волшебник','9 круг','- Да восстанет город Прентон! Да откроются врата Прентона и впустят меня! Да найду я библиотеку вечности в златом городе! Да познаю я секреты всемогущества!']]
            pass
        elif choiceClass==5:# 5. Друид

            pass
        elif choiceClass==6:# 6. Жрец
            # Font Zapf ChanceC
            # 36
            # 8
            TitleCard= [['Жрец','Заговоры','Что мы говорим смерти? Не сегодня – сказал жрец, накладывая на очередного воина заоговор «Уход за умирающим».'],
                        ['Жрец','1 круг',  'Услышь мое слово и встань! Твой бой еще не окончен!'],
                        ['Жрец','2 круг',  'О Тор-Громовержец, дай мне силу Мьельнира, твоего друга и верного боевого товарища, дабы мы могли отправить больше горячих душ на пир в Вальхаллу!'],
                        ['Жрец','3 круг',  'Вы посмели окружить меня, друзей моих и всех людей, что мы поклялись защищать?/nЭто не вы заперли нас. Это мы дадим духам отмщения упиться вашей крови!'],
                        ['Жрец','4 круг',  'И помни, что твоя первая смерть на арене – это всего лишь маленькая смерть. Встань и заверши бой в нашу пользу.'],	
                        ['Жрец','5 круг',  'С неба обрушился невероятно яркий луч. Ударив прямо в лича он моментально испепелил его одежды, а скелеты-прислужники испарились вовсе. '],
                        ['Жрец','6 круг',  'А теперь, дамы и господа, Я предлагаю обняться и смотаться нафиг из этого логова дракона в святилище моего бога!!!'],
                        ['Жрец','7 круг',  '- Нам точно надо тратить этот бриллиант?/n- Да./n- Но ведь я чуть не умер , доставая его!/n- Ты страдал не зря./n- Но…/n- Помолчи и не мешай. О Шесну, великая воительница, победившая дракона Андуина. Драконья угроза снова нависла над этими землями и я прошук тебя… ВОССТАНЬ!'],
                        ['Жрец','8 круг',  'Встань, страх преодолей, встань, в полный рост, встань, на земле своей и достань рукой до звезд.'],
                        ['Жрец','9 круг',  'Латандер, да дарует сила твоя исцеление страждущим воинам, что сражаются за дело твоё светлое.']]

            pass
        elif choiceClass==7:# 7. Изобретатель

            pass
        elif choiceClass==8:# 8. Колдун

            pass
        elif choiceClass==9:# 9. Монах

            pass
        elif choiceClass==10:# 10. Паладин

            pass
        elif choiceClass==11:# 11. Плут

            pass
        elif choiceClass==12:# 12. Следопыт

            pass
        elif choiceClass==13:# 13. Чародей
            pass
        elif choiceClass==14:# 14. Все классы
            pass
        elif choiceClass==15:# 15. Настройки
            print("Настройка левой границы")
            Main_Indent=int(input("Размеры левого отступа(мм): "))
            if Main_Indent<3:
                Main_Indent=3
                print("Размеры левого отступа установлен на 3")
            elif Main_Indent>13:
                Main_Indent=13
                print("Размеры левого отступа установлен на 13")

            columns_in_table=int(input("Колличество колонок: "))
            if columns_in_table<1:
                columns_in_table=1
                print("Количество колонок установлено на 1")
            elif columns_in_table>6:
                columns_in_table=6
                print("Количество колонок установлено на 6")

            rows_in_table=int(input("Колличество строк: "))
            if rows_in_table<1:
                rows_in_table=1
                print("Количество строк установлено на 1")
            elif rows_in_table>3:
                rows_in_table=3
                print("Количество строк установлено на 3")
            # print(Main_Indent+'мм')
            fChoice=True
            pass        
        else:
            print("Нет такого класса")
            fChoice=True
    except:
        print("Нет такого класса")
        fChoice=True
# Main_Indent=13
# Main_Indent=3

f=open('options.txt','w',encoding='utf-8')
f.write(str(Main_Indent) + '\n')
f.write(str(columns_in_table) + '\n')
f.write(str(rows_in_table) + '\n')
f.close()


intTableWidth=intTableWidth/columns_in_table # cm высота места для поля по горизонтали
font_max_size=16-columns_in_table-rows_in_table

for i in range (0,len(listOfClasses)):
            listOfClasses[i]=listOfClasses[i].lower()
listOfLvl=-1
Spell_lvl_s=-1
Spell_lvl_f=-1
while Spell_lvl_s<0 or Spell_lvl_s>10:
    Spell_lvl_s=int(input("Уровень заклинаний, с которого будут создаваться карточки (0-9) или 10 для всех заклинаний: "))
if Spell_lvl_s!=10:
    while Spell_lvl_f<0 or Spell_lvl_f>9 and Spell_lvl_f<Spell_lvl_s:
        Spell_lvl_f=int(input("Уровень заклинаний, до которого будут создаваться карточки: "))
# name_file =listOfClasses[int(choiceClass)-1]+" "+str(listOfLvl)+"_lvl "+ str(datetime.datetime.now().time())

if Spell_lvl_s==10:
    Spell_lvl_s=0
    Spell_lvl_f=9
name_file =listOfClasses[int(choiceClass)-1]+" "+str(Spell_lvl_s)+"-"+str(Spell_lvl_f)+"_lvl"
# else:
#     Spell_lvl_s=listOfLvl
    # Spell_lvl_f=listOfLvl+1


# print(listOfClasses[int(choiceCl
# ass)-1])
kolClass=0
kolArhetip=0
# print("Spisok dlya zapisi")
while listOfSpels[len(listOfSpels)-1]=='':
        listOfSpels.remove('')
# print(listOfSpels)

doc = Document()
# doc.add_heading('Document Title', 0)
section = doc.sections[-1]
section.orientation = 1
# section.page_width = Inches(10/25.4*29.7)
section.page_width =Mm(297)
# section.page_height = Inches(10/25.4*21)
section.page_height = Mm(210) 
section.top_margin=Mm(6)
section.bottom_margin=Mm(6)
section.left_margin = Mm(7)
section.right_margin = Mm(7)

table0 = doc.add_table(rows=(rows_in_table+1)*2+1, cols=columns_in_table)
# modifyBorder(table0)
table0.autofit=False
hdr_cells = table0.rows[0].cells

# cur_row= table0.add_row()
#         cur_row.height_rule = WD_ROW_HEIGHT.EXACTLY
#         cur_row.height=Inches(10/25.4*0.05)

table0.rows[0].height_rule = WD_ROW_HEIGHT.EXACTLY
table0.rows[0].height=Inches(10/25.4*0.1)
for cur_row in range(rows_in_table):
    table0.rows[cur_row+1].height_rule = WD_ROW_HEIGHT.EXACTLY
    table0.rows[cur_row+1].height=Inches(10/25.4*(intTableHeight)/rows_in_table)
# table0.rows[2].height_rule = WD_ROW_HEIGHT.EXACTLY
# table0.rows[2].height=Inches(10/25.4*(intTableHeight)/rows_in_table)
table0.rows[rows_in_table+1].height_rule = WD_ROW_HEIGHT.EXACTLY
table0.rows[rows_in_table+1].height=Inches(10/25.4*0.05)
for cur_row in range(rows_in_table+2,rows_in_table*2+2):
    table0.rows[cur_row].height_rule = WD_ROW_HEIGHT.EXACTLY
    table0.rows[cur_row].height=Inches(10/25.4*(intTableHeight+0.2)/rows_in_table)
# table0.rows[5].height_rule = WD_ROW_HEIGHT.EXACTLY
# table0.rows[5].height=Inches(10/25.4*(intTableHeight+0.2)/rows_in_table)
table0.rows[(rows_in_table+1)*2].height_rule = WD_ROW_HEIGHT.EXACTLY
table0.rows[(rows_in_table+1)*2].height=Inches(10/25.4*0.05)

for columns_in_table_w in range(columns_in_table):
    hdr_cells[columns_in_table_w].width=Inches(10/25.4*intTableWidth)
# hdr_cells[1].width=Inches(10/25.4*intTableWidth)
# hdr_cells[2].width=Inches(10/25.4*intTableWidth)
# hdr_cells[3].width=Inches(10/25.4*intTableWidth)
t=0
row=0
column=0
back_name=listOfClasses[choiceClass-1]
# print (back_name)

for listOfLvl in range (Spell_lvl_s,Spell_lvl_f+1):
    row=(t//8*6)+t//4%2+1
    column=t%4
    hdr_cells = table0.rows[row].cells
    cell = table0.cell(row, column)
    set_cell_border(
        cell,
        left={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
        top={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 1, "val": "single", "color": "#000000", "space": "0"}

    )
    # TitleCard
    if TitleCard[listOfLvl][0]!='':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_after=Mm(0)
        cell.paragraphs[0].paragraph_format.line_spacing=1
        cell.paragraphs[0].paragraph_format.left_indent=Mm(Main_Indent)
        cell.paragraphs[0].paragraph_format.space_before=Mm(1)
        cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cur_run=cell.paragraphs[-1].add_run(TitleCard[listOfLvl][0])
        cur_run.font.name=TitleFont
        cur_run.bold=True
        cur_run.font.size=Pt(26)

        paragraph = cell.add_paragraph('')
        cell.paragraphs[1].paragraph_format.space_after=Mm(0)
        cell.paragraphs[1].paragraph_format.line_spacing=1
        cell.paragraphs[1].paragraph_format.left_indent=Mm(Main_Indent)
        cell.paragraphs[1].paragraph_format.space_before=Mm(1)
        cell.paragraphs[1].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cur_run=paragraph.add_run(TitleCard[listOfLvl][1])
        cur_run.font.name=TitleFont
        cur_run.font.size=Pt(26)

        cell= table0.cell(row+rows_in_table+1,columns_in_table-1-column)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_after=Mm(0)
        cell.paragraphs[0].paragraph_format.line_spacing=1
        cell.paragraphs[0].paragraph_format.left_indent=Mm(Main_Indent)
        cell.paragraphs[0].paragraph_format.space_before=Mm(1)
        cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        cur_run=cell.paragraphs[-1].add_run(TitleCard[listOfLvl][2])
        cur_run.font.name=TitleFont
        cur_run.font.size=Pt(10)

        t=t+1
        add_row(table0,t)
    if listOfLvl==0:
        listOfLvl="Заговор"
    for i in listOfSpels:
        if i[1].find(str(listOfLvl))!=-1 or listOfLvl==10:
            # row=(t//4)
            # // - целочисленное деление; % - Остаток от деления
            # row=t//4+t//8*5
            
            
            
            row=(t//(columns_in_table*rows_in_table)*(rows_in_table+1)*2)+t//(rows_in_table*2)%rows_in_table+1
            # row=(t//8*6)+t//4%2+1
            # print()
            column=t%columns_in_table
            hdr_cells = table0.rows[row].cells
            cell = table0.cell(row, column)
            set_cell_border(
                cell,
                left={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                right={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                top={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                bottom={"sz": 1, "val": "single", "color": "#000000", "space": "0"}

            )
            
            if i[6].find(listOfClasses[3])!=-1 and choiceClass==3 and fFlagChoice==1 :#Мистический рыцарь and (i[1].find('воплощен')!=-1 or i[1].find('огражден')!=-1)
                choiceClass=4
            if i[6].find(listOfClasses[int(choiceClass)-1])!=-1 or i[7].find(listOfClasses[int(choiceClass)-1])!=-1 or choiceClass==14:
                # font_size0=8
                # font_size1=8
                font_size0=font_max_size
                font_size1=font_max_size
                chars_in_line=(intTableWidth-Main_Indent/10)/(font_size0*2/100)
                lenn=0
                cur_String=0
                cur_String_len=0
                for j in range(7,len(i)):
                    lenn=lenn+len(i[j].replace(" ",""))
                    pass
                print(i[0]+" "+str(lenn))
                big_text=False

                cell.paragraphs[0].paragraph_format.space_after=Mm(0)
                cell.paragraphs[0].paragraph_format.line_spacing=1
                cell.paragraphs[0].paragraph_format.left_indent=Mm(Main_Indent)
                cell.paragraphs[0].paragraph_format.space_before=Mm(1)
                # cell.text='!!!'
                # print(cell.paragraphs)
                cur_run=cell.paragraphs[-1].add_run(i[0])
                # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)'
                cur_run.font.name='Cambria (Body)'
                cur_run.bold=True
                cur_run.font.size=Pt(font_max_size)
                cur_run.font.color.rgb=RGBColor(255,255,255)
                cur_run.font.highlight_color = WD_COLOR_INDEX.BLACK
                
                correct=0
                start_txt=9
                # if i[7]=="???":
                #     start_txt=8
                # print(i[2])
                correcting=True
                while correcting:
                    for cur_par in range(1,len(i)):
                        if i[cur_par]!="???":
                            cur_String+=1
                            cur_String_len=0
                            splitted=re.split(' ', i[cur_par])
                            for xi in range(len(splitted)):
                                splitted[xi]=splitted[xi]+" "
                                # cur_run=paragraph.add_run(splitted[xi])
                                if cur_String_len+len(splitted[xi])>chars_in_line:
                                    cur_String_len=len(splitted[xi])
                                    cur_String+=1
                                    # print(splitted[xi])
                                    
                                else:
                                    cur_String_len+=len(splitted[xi])
                    # if cur_String >67-5*font_size0 and font_size0>=5:
                    # print("chars_in_line = "+str(chars_in_line))
                    # print(intTableHeight)
                    # print(font_size0)
                    # print(str((intTableHeight/2-0.7)/((font_size0*5-7)/100)))
                    # print(cur_String)
                    if cur_String > (intTableHeight/rows_in_table-int_Otstup)/((font_size0*5-7)/100) and font_size0>=5:
                         
                        font_size0-=1
                        font_size1-=1
                        # if font_size0==7:
                        if font_size0>=5:
                            chars_in_line=(intTableWidth-Main_Indent/10)/(font_size0*2/100)
                        cur_String=0
                        cur_String_len=0

                    else:
                        correcting=False
                        # print(cur_String)
                    if font_size0==4:
                        big_text=True
                        # print("font_size=4")
                        # Если Big_Text
                        
                paragraph = cell.add_paragraph('')
                cur_run=paragraph.add_run(i[1])
                cur_run.font.name='Cambria (Body)'
                cur_run.font.size=Pt(font_size0)
                cur_run.italic=True
                if not big_text:
                    for cur_par in range(2,start_txt):
                        if i[cur_par]!="???":
                                
                                paragraph = cell.add_paragraph('')
                                cur_run_dvoetoch=i[cur_par].find(":")
                                cur_run=paragraph.add_run(i[cur_par][0:cur_run_dvoetoch+1])
                                
                                cur_run.font.name='Cambria (Body)'
                                cur_run.bold=True
                                cur_run.font.size=Pt(font_size0)
                                cur_run=paragraph.add_run(i[cur_par][cur_run_dvoetoch+1:len(i[cur_par])])
                                
                                cur_run.font.name='Cambria (Body)'
                                cur_run.font.size=Pt(font_size0)

                        else:
                            # print(cur_par)
                            correct=1
                    
                    # Пишем на карточку текст описания заклинания
                    
                    for cur_par in range(start_txt,len(i)):
                        # splitted=i[cur_par].split()
                        
                        paragraph = cell.add_paragraph('')
                        
                        if i[cur_par].find("На больших уровнях")!=-1:
                            # print("На больших уровнях2")
                            cur_run=paragraph.add_run("На больших уровнях")
                            # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)'

                            # for ab in range(len(paragraph.runs),0,-1):
                            #     paragraph.runs[ab]=paragraph.runs[ab-1]
                        
                            cur_run.font.name='Cambria (Body)'
                            cur_run.font.size=Pt(font_size1)
                            cur_run.bold=True
                            i[cur_par]=i[cur_par][i[cur_par].find("На больших уровнях")+len("На больших уровнях"):len(i[cur_par])]
                            # print(i[cur_par])
                        elif i[cur_par].find("На высоких уровнях")!=-1:
                            cur_run=paragraph.add_run("На высоких уровнях")
                            # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)'

                            # for ab in range(len(paragraph.runs),0,-1):
                            #     paragraph.runs[ab]=paragraph.runs[ab-1]
                        
                            cur_run.font.name='Cambria (Body)'
                            cur_run.font.size=Pt(font_size1)
                            cur_run.bold=True
                            i[cur_par]=i[cur_par][i[cur_par].find("На высоких уровнях")+len("На высоких уровнях"):len(i[cur_par])]
                            # print(i[cur_par])

                        if i[cur_par].find("Компонент авторских отчислений (А)")!=-1:
                            # print("Авторские отчисления")
                            cur_run=paragraph.add_run("Компонент авторских отчислений (А) ")
                            # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)'
                            cur_run.font.name='Cambria (Body)'
                            cur_run.font.size=Pt(font_size1)
                            cur_run.italic=True
                            cur_run.bold=True
                            i[cur_par]=i[cur_par][i[cur_par].find("Компонент авторских отчислений (А)")+len("Компонент авторских отчислений (А)"):len(i[cur_par])]
                        
                        cur_run=paragraph.add_run(i[cur_par])
                        # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)' 
                        cur_run.font.name='Cambria (Body)'
                        cur_run.font.size=Pt(font_size1)
                        cur_run.bold=False
                        for cur_par1 in range(1,len(cell.paragraphs)):
                            cell.paragraphs[cur_par1].paragraph_format.space_after=Mm(0)
                            cell.paragraphs[cur_par1].paragraph_format.line_spacing=1
                            cell.paragraphs[cur_par1].paragraph_format.left_indent=Mm(Main_Indent)
                            if (cell.paragraphs[cur_par1].text.find(":")<21 and cell.paragraphs[cur_par1].text.find(":")>-1 or cur_par1<4) and cell.paragraphs[cur_par1].text.find("На высоких уровнях")==-1:
                                cell.paragraphs[cur_par1].paragraph_format.first_line_indent=Mm(0)
                                a=cell.paragraphs[cur_par1].text.find(":")
                                if a>-1:
                                    for ab in range(0,len(cell.paragraphs[cur_par1].runs)):
                                        cell.paragraphs[cur_par1].runs[ab].bold=True
                                        if cell.paragraphs[cur_par1].runs[ab].text.find(":")!=-1:
                                            break
                                    pass
                            elif (cell.paragraphs[cur_par1].text.find(":")<21 and cell.paragraphs[cur_par1].text.find(":")>-1 or cur_par1<4) and cell.paragraphs[cur_par1].text.find("На больших уровнях")==-1:
                                cell.paragraphs[cur_par1].paragraph_format.first_line_indent=Mm(0)
                                a=cell.paragraphs[cur_par1].text.find(":")
                                if a>-1:
                                    for ab in range(0,len(cell.paragraphs[cur_par1].runs)):
                                        cell.paragraphs[cur_par1].runs[ab].bold=True
                                        if cell.paragraphs[cur_par1].runs[ab].text.find(":")!=-1:
                                            break
                                    pass
                            
                            else:
                                cell.paragraphs[cur_par1].paragraph_format.first_line_indent=Mm(1)
                        for ab in range(0,len(cell.paragraphs[1].runs)):
                            cell.paragraphs[1].runs[ab].bold=False
                        


                else:
                    # Пишем большой текст с учётом обеих сторон
                    # Вычисляем размер шрифта для большого текста
                    cur_String=0
                    cur_String_len=0
                    font_size1=1
                    # print("Big Text!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
                    font_size0=font_max_size
                    
                    while font_size0>4:
                    # for font_size0 in range(8,4,-1) :
                        chars_in_line=(intTableWidth-Main_Indent/10)/(font_size0*2/100)
                        for cur_par in range(1,len(i)):
                            if i[cur_par]!="???":
                                cur_String+=1
                                cur_String_len=0
                                splitted=re.split(' ', i[cur_par])
                                for xi in range(len(splitted)):
                                        splitted[xi]=splitted[xi]+" "
                                        # cur_run=paragraph.add_run(splitted[xi])
                                        if cur_String_len+len(splitted[xi])>chars_in_line:
                                            cur_String_len=len(splitted[xi])
                                            cur_String+=1
                                            # print(splitted[xi])
                                            
                                        else:
                                            cur_String_len+=len(splitted[xi])
                                            
                            # if cur_String >(67-5*font_size0)*2 and font_size0>5:
                        if cur_String >(intTableHeight/rows_in_table-int_Otstup)/((font_size0*5-7)/100)*2 and font_size0>5:
                                cur_String=0
                                cur_String_len=0
                                font_size0-=1
                        else:
                                font_size1=font_size0
                                pass
                        #     if  font_size1!=1:
                        #         break
                        if  font_size1!=1:
                                break    
                        
                    font_size0=font_size1
                    # print("Font = "+str(font_size0))

                    # Вписываем текст после вычисления размера шрифта для большого текста
                    cur_String=0
                    cur_String_len=0
                    bSecondSide=False
                    cell.paragraphs[1].paragraph_format.space_after=Mm(0)
                    cell.paragraphs[1].paragraph_format.line_spacing=1
                    cell.paragraphs[1].paragraph_format.left_indent=Mm(Main_Indent)
                    cell.paragraphs[1].paragraph_format.first_line_indent=Mm(0)
                    # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)' 
                    cur_run.font.name='Cambria (Body)'
                    cur_run.font.size=Pt(font_size1)
                    paragraph = cell.add_paragraph('') 
                    paragraph.paragraph_format.left_indent=Mm(Main_Indent)
                    paragraph.paragraph_format.first_line_indent=Mm(0) 
                    correct=1
                    
                    for cur_par in range(2,len(i)):
                        
                        if i[cur_par]!="???":
                            cur_String+=1
                            if cur_String >(intTableHeight/rows_in_table-int_Otstup)/((font_size0*5-7)/100):
                            # if cur_String >(intTableHeight/2-0.05)/((font_size0*5-7)/100):
                            # if cur_String >(67-5*font_size0):
                                cell=jump_to_cell(cell)
                                bSecondSide=True
                                paragraph = cell.paragraphs[0] 
                                cell.paragraphs[0].paragraph_format.space_before=Mm(3)
                                cell.paragraphs[0].paragraph_format.space_after=Mm(0)
                                cell.paragraphs[0].paragraph_format.line_spacing=1
                                cell.paragraphs[0].paragraph_format.right_indent=Mm(Main_Indent)
                                cell.paragraphs[0].paragraph_format.left_indent=Mm(6)
                                cell.paragraphs[0].paragraph_format.first_line_indent=Mm(0)
                                # paragraph = cell.add_paragraph('') 
                                cur_String=0
                                cur_String_len=len(splitted[xi])
                            cur_String_len=0
                            splitted=re.split(' ', i[cur_par])
                            for xi in range(len(splitted)):
                                splitted[xi]=splitted[xi]+" "
                                
                                if cur_String_len+len(splitted[xi])>chars_in_line:
                                    cur_String_len=len(splitted[xi])
                                    cur_String+=1
                                    if cur_String >(intTableHeight/rows_in_table-int_Otstup)/((font_size0*5-7)/100):
                                        # print(cur_String)
                                        # break
                                        cell=jump_to_cell(cell)
                                        bSecondSide=True
                                        paragraph = cell.paragraphs[0] 
                                        cell.paragraphs[0].paragraph_format.space_before=Mm(3)
                                        cell.paragraphs[0].paragraph_format.space_after=Mm(0)
                                        cell.paragraphs[0].paragraph_format.line_spacing=1
                                        cell.paragraphs[0].paragraph_format.right_indent=Mm(Main_Indent)
                                        cell.paragraphs[0].paragraph_format.left_indent=Mm(6)
                                        cell.paragraphs[0].paragraph_format.first_line_indent=Mm(0)
                                        # paragraph = cell.add_paragraph('') 
                                        cur_String=0
                                        cur_String_len=len(splitted[xi])
                                        pass
                                else:
                                    cur_String_len+=len(splitted[xi])
                                cur_run=paragraph.add_run(splitted[xi])
                                # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)'
                                cur_run.font.name='Cambria (Body)'
                                cur_run.font.size=Pt(font_size1)

                                    # print(splitted[xi])
                                        
                                    
                            
                            paragraph = cell.add_paragraph('')  
                    for cur_par in range(1,len(cell.paragraphs)):
                        cell.paragraphs[cur_par].paragraph_format.space_after=Mm(0)
                        cell.paragraphs[cur_par].paragraph_format.line_spacing=1
                        # cell.paragraphs[cur_par].paragraph_format.left_indent=Mm(Main_Indent)
                        cell.paragraphs[cur_par].paragraph_format.right_indent=Mm(Main_Indent)
                        cell.paragraphs[cur_par].paragraph_format.left_indent=Mm(6)
                        cell.paragraphs[cur_par].paragraph_format.first_line_indent=Mm(1)  
                        # cell.paragraphs[cur_par].paragraph_format.first_line_indent=Mm(0)
                            
                    pass

# ------------------------------------------------------------------------------------------------------------------------------------------------------------------------
                
                # Иначе большими буквами название заклинания, круг и классы
                if not big_text:        
                    cell= table0.cell(row+rows_in_table+1,columns_in_table-1-column)
                    # cell= table0.cell(row+3, 3-column)
                    cell.paragraphs[0].paragraph_format.right_indent=Mm(Main_Indent)
                    cell.paragraphs[0].paragraph_format.left_indent=Mm(8)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.paragraphs[0].paragraph_format.right_indent=Mm(Main_Indent)
                    cell.paragraphs[0].paragraph_format.left_indent=Mm(5)
                    txt=i[0]
                    # txt=txt[0:txt.find("[")-1]
                    cur_run=cell.paragraphs[0].add_run(txt)
                    # # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)'; cur_run.font.name='Cambria (Body)'
                    cur_run.font.size=Pt(font_max_size+8)
                    cur_run.bold=True
                    paragraph = cell.add_paragraph('')
                    paragraph.paragraph_format.right_indent=Mm(Main_Indent)
                    paragraph.paragraph_format.left_indent=Mm(8)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if listOfLvl==0 or listOfLvl=="Заговор":
                        cur_run=paragraph.add_run("Заговор")
                        # # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)'; cur_run.font.name='Cambria (Body)'
                    else:

                        cur_run=paragraph.add_run(str(listOfLvl) + " круг")
                        # # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)'; cur_run.font.name='Cambria (Body)'
                    cur_run.font.size=Pt(7+font_max_size)
                    txt=i[6]
                    txt=txt[txt.find(":")+2:len(txt)]
                    txt=txt.title()
                    paragraph = cell.add_paragraph('')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.right_indent=Mm(Main_Indent)
                    paragraph.paragraph_format.left_indent=Mm(5)
                    paragraph.paragraph_format.line_spacing=1
                    cur_run=paragraph.add_run(txt)
                    cur_run.font.size=Pt(0+font_max_size)
                    if i[7]!="???":
                        txt=i[7]
                        txt=txt[txt.find(":")+2:len(txt)]
                        # paragraph = cell.add_paragraph('')
                        # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cur_run=paragraph.add_run(", "+txt)
                        # # cur_run.font.name='TeXGyreCursor'; cur_run.font.name='Cambria (Body)'; cur_run.font.name='Cambria (Body)'
                        cur_run.font.size=Pt(0+font_max_size)
                    # cur_run.font.all_caps=True

                if i[6].find(listOfClasses[int(choiceClass)-1])!=-1:
                    kolClass+=1
                if i[7].find(listOfClasses[int(choiceClass)-1])!=-1:
                    # print(i[7])
                    kolArhetip+=1
                # print(i[8])
                
                # print()
                # toDoc(i,doc,t)
                # print ("t="+str(t)+"  row="+str(row))
                t=t+1
                add_row(table0,t)
            # if t>=16:
            #     break
            

                
                
            
        
    
print("Klass: "+str(kolClass)+"  Arhetips: "+str(kolArhetip))
# name =listOfClasses[int(choiceClass)-1]+" "+str(listOfLvl)+"_lvl "+ str(datetime.datetime.now().time())
name_file=name_file.replace(':','_',2)
try:
    os.mkdir("Docs")
except:
    pass
doc.save('Docs/'+name_file+'.docx')
convert('Docs/'+name_file+'.docx')